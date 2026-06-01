import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# Create a new document
doc = docx.Document()

# Set margins to 1 inch (72 points)
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Define colors
dark_color = RGBColor(34, 34, 34)  # Hex #222222
gray_color = RGBColor(102, 102, 102)  # Hex #666666

def set_font(run, size=11, bold=False, italic=False, color=dark_color):
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    # Fix for some Word versions ignoring font.name
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

def add_para(text, size=11, bold=False, italic=False, color=dark_color, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=5, line_spacing=1.15):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    run = p.add_run(text)
    set_font(run, size, bold, italic, color)
    return p

def add_heading(text):
    return add_para(text, size=12, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=4, line_spacing=1.0)

def add_bullet_bold_prefix(bold_text, normal_text):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    
    run1 = p.add_run(bold_text)
    set_font(run1, size=11, bold=True)
    
    run2 = p.add_run(normal_text)
    set_font(run2, size=11, bold=False)
    return p

def add_body_bold_prefix(bold_text, normal_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    
    run1 = p.add_run(bold_text)
    set_font(run1, size=11, bold=True)
    
    run2 = p.add_run(normal_text)
    set_font(run2, size=11, bold=False)
    return p

# TITLE
add_para("RADAR", size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3, line_spacing=1.0)
add_para("Risk Analysis Data Adaptive Response", size=11, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, line_spacing=1.0)
add_para("R&D Techathon 2026 — Proje Başvurusu", size=10, color=gray_color, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18, line_spacing=1.0)

# SORU 1
add_heading("Projede yapay zekâ nerede kullanılıyor?")
add_para("Proje, kurumun DLP (Data Loss Prevention) altyapısından toplanan kullanıcı davranış verilerini yapay zekâ ile analiz ederek dinamik risk skorlaması yapan ve koruma politikalarını otomatik uyarlayan bir platformdur. Yapay zekâ şu noktalarda kullanılır:")
add_bullet_bold_prefix("Davranışsal risk analizi: ", "Kullanıcıların dosya hareketleri, e-posta gönderim örüntüleri, USB kullanımı, ekran görüntüsü alma gibi DLP olaylarını analiz ederek her kullanıcıya dinamik bir risk skoru (0–100) atanması.")
add_bullet_bold_prefix("Anomali tespiti: ", "Kullanıcının 20 günlük kayar pencere (rolling window) bazlı geçmiş davranış profilinden sapmaların 3-sigma kuralı ile istatistiksel olarak algılanması; ani yükselişler, olağandışı kanal kullanımı, yeni hedef adreslere veri çıkışı gibi anomalilerin otomatik tespiti.")
add_bullet_bold_prefix("AI destekli davranış açıklaması: ", "Azure OpenAI entegrasyonu ile her kullanıcının risk profilinin doğal dilde açıklanması — “Bu kullanıcı son 7 günde normal ortalamanın 3 katı hassas dosya transferi gerçekleştirdi” gibi anlaşılır özetlerin otomatik üretilmesi.")
add_bullet_bold_prefix("Akıllı sınıflandırma: ", "DLP olaylarının hassasiyet düzeyine göre otomatik sınıflandırılması (düşük / orta / yüksek / kritik).")
add_bullet_bold_prefix("Copilot (ChatBot) arayüzü: ", "Güvenlik analistlerinin doğal dilde sorular sorarak (“En riskli 5 kullanıcı kimdir?”, “Geçen hafta hangi departmandan en çok ihlal geldi?”) sistemi sorgulamasını sağlayan AI chatbot.")

# SORU 2
add_heading("Projenin konusu ve kapsamı nedir?")
add_para("Kurum, Forcepoint DLP altyapısı üzerinden veri kaybı önleme politikaları uyguluyor. Ancak mevcut DLP sistemi statik, kural tabanlı çalışıyor: her kullanıcıya aynı politika aynı şiddette uygulanıyor. Bir kullanıcının düşük riskli bir dosya paylaşımıyla, kasıtlı veri sızdırma girişimi aynı seviyede değerlendirildiğinde hem güvenlik ekipleri gereksiz uyarı yığını altında kalıyor hem de gerçek tehditler gözden kaçabiliyor.")
add_body_bold_prefix("RADAR", ", kurumun DLP altyapısından akan olay verilerini toplayarak kullanıcı bazlı risk skorlaması yapan, bu skora göre DLP politikalarını dinamik olarak uyarlayan ve güvenlik ekiplerine AI destekli içgörüler sunan bir platformdur. Proje kapsamı:")
add_bullet_bold_prefix("Veri Toplama (Collector): ", "Forcepoint DLP’den Splunk/SIEM üzerinden akan olay loglarının gerçek zamanlı toplanması ve standart formata dönüştürülmesi.")
add_bullet_bold_prefix("Risk Analizi (Analyzer): ", "Toplanan verilerin yapay zekâ algoritmaları ile analiz edilmesi; kullanıcı bazlı risk skoru, trend analizi ve anomali tespiti.")
add_bullet_bold_prefix("Davranış Motoru (Behavior Engine): ", "Kullanıcının günlük, haftalık ve aylık davranış profilinin Z-score tabanlı çok boyutlu analiz ile çıkarılması; 30/60/90 gün adaptif baseline seçimi; departman ve rol bazlı karşılaştırmalar; IOB (Indicators of Behavior) tespiti.")
add_bullet_bold_prefix("Adaptif Politika Yönetimi: ", "Risk skoruna göre DLP politikalarının otomatik sıkılaştırılması veya gevşetilmesi (izin ver → uyar → onayla → engelle).")
add_bullet_bold_prefix("Dashboard ve Raporlama: ", "Gerçek zamanlı risk haritası, kullanıcı detay sayfaları, trend grafikleri ve PDF/Excel rapor üretimi.")
add_bullet_bold_prefix("Olay Yönetimi (Investigation): ", "Yüksek riskli olayların incelenmesi, remediation (düzeltme) aksiyonlarının takibi ve denetim izi (audit trail) oluşturulması.")
add_para("Proje, mevcut DLP altyapısının yerine geçmez; ona paralel çalışan bir akıl katmanı olarak konumlanır.")

# SORU 3
add_heading("Projenin amacı ve hedefi nedir? (Asıl fayda kime?)")
add_body_bold_prefix("Asıl fayda kuruma ve dolaylı olarak müşteriyedir. ", "Projenin temel yaklaşımı şudur: güvenlik, iş yapabilirliğin önündeki engel değil, onu mümkün kılan zemin olmalıdır.")
add_para("Klasik DLP yaklaşımı “herkese aynı katılığı uygula” mantığıyla çalışır. Bu, düşük riskli çalışanların gereksiz yere engellenmesine (verimlilik kaybı) ve yüksek riskli davranışların alarm yığını içinde kaybolmasına (güvenlik açığı) neden olur. RADAR bu dengeyi kurumun lehine çevirir:")
add_bullet_bold_prefix("Güvenlik ekiplerinin gerçek tehditlere odaklanması ", "— yanlış pozitif oranının %60–80 azaltılması hedefi.")
add_bullet_bold_prefix("", "Düşük riskli kullanıcılara gereksiz engel konmaması — çalışan memnuniyeti ve operasyonel verimlilik artışı.")
add_bullet_bold_prefix("", "Yüksek riskli davranışların proaktif tespiti — veri sızıntısı gerçekleşmeden müdahale imkânı.")
add_bullet_bold_prefix("", "İç tehditlerin (insider threat) davranış analizi ile erken dönemde yakalanması.")
add_bullet_bold_prefix("", "KVKK ve düzenleyici kurum gerekliliklerine uyumlu denetim izinin otomatik oluşturulması.")
add_bullet_bold_prefix("", "Kurum içi güvenliğin güçlenmesi, müşteri verisinin daha iyi korunması anlamına gelir.")
add_bullet_bold_prefix("", "Hassas müşteri verisine erişen kullanıcıların risk skorları sürekli izlenir; anormal davranış tespit edildiğinde otomatik politika sıkılaştırması ile müşteri verisi korunur.")

# SORU 4
add_heading("Mevcut durumda sorun ne?")
add_bullet_bold_prefix("Statik politika sorunu: ", "Forcepoint DLP, kural tabanlı çalışıyor. “Hassas dosya USB’ye kopyalanırsa engelle” kuralı, finans departmanındaki bir çalışan ile IT destek personeli için aynı şekilde uygulanıyor. Bağlam yok, risk değerlendirmesi yok.")
add_bullet_bold_prefix("Uyarı yorgunluğu (Alert Fatigue): ", "DLP sistemi günde yüzlerce/binlerce olay üretiyor. Güvenlik ekibi bu olayları manuel olarak inceliyor; çoğu düşük riskli veya yanlış pozitif. Gerçek tehditler uyarı yığınının içinde kayboluyor.")
add_bullet_bold_prefix("Davranış bağlamı eksikliği: ", "Mevcut DLP “ne oldu” sorusuna cevap veriyor ama “kim, ne sıklıkla, normalden ne kadar farklı” sorularına cevap veremiyor. Bir kullanıcının ilk kez mi yoksa yüzüncü kez mi ihlal yaptığı bilinmiyor.")
add_bullet_bold_prefix("Reaktif yaklaşım: ", "Mevcut sistem olay gerçekleştikten sonra log tutuyor. Proaktif risk tespiti, trend analizi ve erken uyarı mekanizması bulunmuyor.")
add_bullet_bold_prefix("Merkezi görünürlük eksikliği: ", "Kullanıcı bazlı bütünleşik bir risk görünümü yok. Farklı kanallardan (e-posta, endpoint, ağ) gelen olaylar ayrı ayrı değerlendiriliyor; kullanıcının toplam risk profili oluşturulamıyor.")
add_bullet_bold_prefix("Manuel politika yönetimi: ", "Politika değişiklikleri güvenlik ekibi tarafından manuel yapılıyor. Riskin dinamik doğasına ayak uydurmak pratik olarak mümkün değil.")

# SORU 5
add_heading("Rakiplerden farkı ne?")
add_para("Bu alanda güçlü oyuncular var. Aşağıda temel karşılaştırma:")
add_body_bold_prefix("Forcepoint Risk-Adaptive Protection (RAP): ", "130+ davranış göstergesi (IOB) ile kullanıcı risk skoru (0–100) hesaplıyor; DLP politikalarını otomatik uyarlıyor. ARIA (AI asistan) ile doğal dilde politika oluşturma desteği sunuyor. Ancak ticari lisans maliyeti çok yüksek, kuruma özel uyarlama imkânı sınırlı (vendor lock-in), Türkçe ve KVKK’ya özel uyarlama yok.")
add_body_bold_prefix("Microsoft Purview Adaptive Protection: ", "Insider Risk Management ile DLP’yi entegre ederek kullanıcılara Minor/Moderate/Elevated risk seviyesi atıyor. Ancak yalnızca Microsoft ekosisteminde (Exchange, Teams, Endpoint) çalışıyor. Forcepoint DLP gibi üçüncü parti DLP verileriyle entegre olmuyor. E5 lisansı gerektiriyor.")
add_body_bold_prefix("Symantec / Digital Guardian / Trellix DLP: ", "UEBA (User Entity Behavior Analytics) modülleri ile risk skorlama yapıyorlar. Ancak risk-adaptif politika uyarlama bu ürünlerde birden fazla platform entegrasyonu gerektiriyor; tek başlarına tam adaptif döngü sunmuyorlar.")
add_body_bold_prefix("Bizim farkımız: ", "(1) Aynı vizyonu, kurumun kendi altyapısında, açık kaynak teknolojiler üzerine kurulu olarak sunuyoruz — lisans maliyeti sıfır. (2) Vendor-agnostik: Forcepoint, Splunk veya herhangi bir SIEM/DLP kaynağından veri alabiliriz. (3) Uçtan uca döngü: veri toplama → analiz → risk skoru → politika uyarlama → remediation → doğrulama. (4) Türkçe, KVKK ve kurum bağlamına özel. Bu birleşimi sunan bir örnek yok.")

# SORU 6
add_heading("Proje hangi adımlarla uygulanır?")
add_bullet_bold_prefix("", "Veri toplama altyapısı kurulumu: Forcepoint DLP / Splunk’tan olay loglarının Redis Stream üzerinden gerçek zamanlı toplanması; veri formatının standartlaştırılması ve veritabanına (PostgreSQL) yazılması.")
add_bullet_bold_prefix("", "Risk skorlama motorunun geliştirilmesi: Kullanıcı bazlı çok boyutlu risk skoru algoritmasının oluşturulması — olay türü ağırlıkları, sıklık analizi, hassasiyet seviyesi, zaman bağlamı, departman bazlı normal davranış profili.")
add_bullet_bold_prefix("", "Anomali tespit modülü: Kullanıcının geçmiş davranış ortalamasından sapmaların istatistiksel ve yapay zekâ tabanlı yöntemlerle tespit edilmesi; anlık risk artışının tetiklenmesi.")
add_bullet_bold_prefix("", "AI davranış analizi entegrasyonu: Azure OpenAI / OpenAI API ile davranış profillerinin doğal dilde açıklanması; güvenlik analistleri için anlaşılır, eyleme dönüştürülebilir özetlerin üretilmesi.")
add_bullet_bold_prefix("", "Adaptif politika yönetimi: Risk skoru seviyelerine göre DLP politikalarının otomatik uyarlanması — düşük risk: izin ver, orta risk: uyarı göster, yüksek risk: onay iste, kritik risk: engelle.")
add_bullet_bold_prefix("", "Dashboard ve raporlama: Next.js tabanlı modern web arayüzünde gerçek zamanlı risk haritası, kullanıcı detay sayfaları, trend grafikleri, olay inceleme ekranları ve otomatik rapor üretimi (PDF/Excel).")
add_bullet_bold_prefix("", "Olay yönetimi ve remediation: Yüksek riskli olayların inceleme sürecine alınması, düzeltme aksiyonlarının takibi ve denetim izinin tutulması.")
add_bullet_bold_prefix("", "Test ve doğrulama: DLP test senaryoları ile sistemin uçtan uca doğrulanması; farklı risk seviyelerinde politika uyarlama davranışının test edilmesi.")
add_bullet_bold_prefix("", "Pilot uygulama: Seçili bir departman veya kullanıcı grubu üzerinde kontrollü pilot çalışma; sonuçların ölçülmesi ve ince ayar.")

# SORU 7
add_heading("Hangi teknolojiler kullanılacak?")
add_bullet_bold_prefix("Backend API: ", "ASP.NET Core 8 (C#) — risk analizi, politika yönetimi, kullanıcı yönetimi REST API.")
add_bullet_bold_prefix("Veri Toplama: ", "Redis Streams — Forcepoint/Splunk’tan gelen olay verilerinin gerçek zamanlı işlenmesi.")
add_bullet_bold_prefix("Veritabanı: ", "PostgreSQL — kullanıcı profilleri, risk skorları, olay logları, konfigürasyon.")
add_bullet_bold_prefix("AI / LLM: ", "Azure OpenAI / OpenAI API — davranış açıklaması, chatbot, anomali yorumlama.")
add_bullet_bold_prefix("Web Dashboard: ", "Next.js (React) + TypeScript + TailwindCSS — modern, gerçek zamanlı yönetim arayüzü.")
add_bullet_bold_prefix("Raporlama: ", "ClosedXML (Excel) + QuestPDF (PDF) — otomatik rapor üretimi.")
add_bullet_bold_prefix("Arka Plan: ", "Background Services (.NET) — sürekli çalışan analiz, senkronizasyon ve temizlik görevleri.")
add_bullet_bold_prefix("SIEM Entegrasyon: ", "Splunk REST API — DLP olay verilerinin çekilmesi.")
add_bullet_bold_prefix("Container: ", "Docker — dağıtım ve ortam bağımsızlığı.")
add_bullet_bold_prefix("CI/CD: ", "GitHub Actions — otomatik build, test ve dağıtım pipeline’ı.")

# SORU 8
add_heading("Riskler, kısıtlar ve maliyetler neler?")
add_body_bold_prefix("Veri kalitesi riski: ", "DLP sisteminden gelen olay verilerinin formatı, eksiksizliği ve tutarlılığı risk skorlamasının doğruluğunu doğrudan etkiler. Önlem: Veri doğrulama ve temizleme katmanı; eksik/tutarsız verilerin raporlanması.")
add_body_bold_prefix("Yanlış pozitif/negatif riski: ", "Risk skoru algoritması ilk aşamada yeterli hassasiyete ulaşamayabilir. Önlem: İnsan denetimi destekli çalışma; pilot dönemde “yalnızca izle” modunda başlayıp kademeli olarak otomatik politika uyarlamaya geçiş.")
add_body_bold_prefix("Adaptif politika yan etkileri: ", "Otomatik politika değişikliklerinin öngörülemeyen operasyonel etkileri olabilir. Önlem: Politika değişikliklerinde sandbox modu; değişiklik öncesi denetim izi ve geri alma mekanizması.")
add_body_bold_prefix("Kısıt: ", "Proje, mevcut DLP altyapısının (Forcepoint) ürettiği olay verilerine bağımlıdır. DLP politikalarının kendisine müdahale etmez; adaptif uyarlama, DLP API’si üzerinden gerçekleşir.")
add_body_bold_prefix("Maliyet: ", "Açık kaynak teknolojiler (ASP.NET Core, PostgreSQL, Next.js, Redis) üzerine kuruludur; lisans maliyeti yoktur. Ana maliyet geliştirme süresi ve Azure OpenAI API kullanım ücretidir (token bazlı, düşük–orta düzey). Mevcut kurum altyapısı kullanılabilir.")

# SORU 9
add_heading("Finansal faydası ne?")
add_para("Güvenlik araçları ilk bakışta yalnızca bir maliyet kalemi gibi görünür; ancak bu proje doğrudan gelire dönüşen temel faydalar sağlar:")
add_bullet_bold_prefix("Ticari ürün lisans tasarrufu: ", "Forcepoint RAP veya Microsoft Purview E5 gibi ticari risk-adaptif çözümlerin yıllık lisans maliyeti kullanıcı başına 30–80 USD arasındadır. Kurum ölçeğinde yıllık yüz binlerce USD ticari lisans maliyetinden kaçınılır.")
add_bullet_bold_prefix("Güvenlik ekibi verimlilik artışı: ", "Yanlış pozitif azaltımı ile güvenlik analistlerinin olay inceleme süresinin %40–60 düşürülmesi hedeflenir.")
add_bullet_bold_prefix("Veri sızıntısı önleme: ", "IBM’in 2025 Cost of a Data Breach raporuna göre bir veri ihlalinin ortalama maliyeti 4,88 milyon USD’dir. Proaktif risk tespiti ile bu riskin ciddi ölçüde azaltılması hedeflenir.")
add_bullet_bold_prefix("Düzenleyici ceza riski azaltımı: ", "KVKK kapsamında veri ihlali durumunda 1–10 milyon TL arası idari para cezası uygulanabilir. Denetim izi ve uyum raporlaması ile ceza riskinin minimize edilmesi.")
add_bullet_bold_prefix("Operasyonel verimlilik: ", "Düşük riskli kullanıcıların gereksiz engellerle karşılaşmaması, iş süreçlerinin aksamadan devam etmesi.")
add_bullet_bold_prefix("Ölçeklenebilirlik: ", "Tek bir araçla kurumdaki tüm DLP kullanıcılarının risk profilinin yönetilmesi — yüksek tekrar kullanım değeri. İleride farklı veri kaynakları (e-posta güvenliği, CASB, endpoint) eklenerek kapsamın genişletilebilmesi.")

# Save document
doc.save(r"c:\Users\abdul\Desktop\dlp-risk-adaptive-protection-csharp-main\Ideathon_Basvuru\RADAR_Basvuru_Final.docx")
print("Done!")
